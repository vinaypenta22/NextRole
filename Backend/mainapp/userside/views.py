from rest_framework.views import APIView
from .serializer import userSerializer
from .serializer import AppliedJobSerializer, SavedJobSerializer, ChatMessageSerializer
from rest_framework.response import Response
from rest_framework import status
from django.contrib.auth.hashers import check_password
from rest_framework_simplejwt.tokens import RefreshToken
from .models import userModel, Resume, Job, AppliedJob, SavedJob, ChatMessage
from .resume_ai_service import ResumeAIService
from docx import Document
import fitz
import os
import json
import re
import time
try:
    from groq import Groq
except Exception:
    Groq = None
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    load_dotenv = None
try:
    import requests
except Exception:
    requests = None

try:
    import serpapi
except Exception:
    serpapi = None
client = Groq(api_key=os.getenv("GROQ_API_KEY")) if Groq else None
class userAuthAPIView(APIView):

    def post(self, request):

        email = request.data.get("email")

        if userModel.objects.filter(email=email).exists():
            return Response(
                {
                    "message": "User already exists."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        serializer = userSerializer(data=request.data)

        if serializer.is_valid():
            serializer.save()
            return Response(
                {
                    "message": "User created successfully.",
                    "data": serializer.data
                },
                status=status.HTTP_201_CREATED
            )

        return Response(
            serializer.errors,
            status=status.HTTP_400_BAD_REQUEST
        )


#login
class UserLoginAPIView(APIView):

    def post(self, request):

        email = request.data.get("email")
        password = request.data.get("password")
        if not email or not password:
            return Response(
                {
                    "message": "Email and Password are required."
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            user = userModel.objects.get(email=email)
        except userModel.DoesNotExist:
            return Response(
                {
                    "message": "Invalid Email or Password."
                },
                status=status.HTTP_401_UNAUTHORIZED,
            )

        if not check_password(password, user.password):
            return Response(
                {
                    "message": "Invalid Email or Password."
                },
                status=status.HTTP_401_UNAUTHORIZED,
            )

        refresh = RefreshToken.for_user(user)

        return Response(
            {
                "message": "Login Successful",
                "user": {
                    "id": user.id,
                    "email": user.email,
                },
                "access": str(refresh.access_token),
                "refresh": str(refresh),
            },
            status=status.HTTP_200_OK,
        )
        
class ResumeUploadAPIView(APIView):

    def post(self, request):

        resume = request.FILES.get("resume")
        user_id = request.data.get("user_id")

        if not resume:
            return Response(
                {
                    "message": "Please upload a resume."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        if not user_id:
            return Response(
                {
                    "message": "user_id is required."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            user = userModel.objects.get(id=user_id)

        except userModel.DoesNotExist:

            return Response(
                {
                    "message": "User not found."
                },
                status=status.HTTP_404_NOT_FOUND
            )

        extension = os.path.splitext(
            resume.name
        )[1].lower()

        if extension not in [".pdf", ".docx"]:

            return Response(
                {
                    "message": "Only PDF and DOCX files are supported."
                },
                status=status.HTTP_400_BAD_REQUEST
            )

        resume_data = Resume.objects.create(
            user=user,
            resume=resume
        )

        try:

            if extension == ".pdf":
                text = self.extract_pdf_text(
                    resume_data.resume.path
                )
            else:
                text = self.extract_docx_text(
                    resume_data.resume.path
                )

            resume_data.extracted_text = text

            extracted_data = self.extract_resume_details(
                text
            )

            if not isinstance(extracted_data, dict):
                extracted_data = {
                    "name": "",
                    "email": "",
                    "phone": "",
                    "skills": [],
                    "experience": "",
                    "experience_years": 0,
                    "education": [],
                    "projects": [],
                    "certifications": [],
                    "current_company": "",
                    "current_designation": "",
                    "location": "",
                    "linkedin": "",
                    "github": "",
                }

            designation = extracted_data.get("current_designation") or extracted_data.get("title") or extracted_data.get("role")
            if isinstance(designation, str) and designation.strip():
                extracted_data["current_designation"] = designation.strip()

            resume_data.is_resume_uploaded = True
            resume_data.resume_details = extracted_data
            resume_data.skills = extracted_data.get(
                "skills",
                []
            )

            resume_data.experience = extracted_data.get(
                "experience",
                ""
            )

            resume_data.education = extracted_data.get(
                "education",
                []
            )

            resume_data.projects = extracted_data.get(
                "projects",
                []
            )

            resume_data.certifications = extracted_data.get(
                "certifications",
                []
            )

            resume_data.save()

            service = ResumeAIService()
            resume_insights = service.generate_ats_score(extracted_data) or {}

            resume_data.last_search_filters = {}
            resume_data.last_recommended_jobs = []
            resume_data.search_query = ""
            resume_data.provider_responses = {}
            resume_data.resume_insights = resume_insights
            resume_data.save(update_fields=[
                "is_resume_uploaded",
                "resume_details",
                "skills",
                "experience",
                "education",
                "projects",
                "certifications",
                "last_search_filters",
                "last_recommended_jobs",
                "search_query",
                "provider_responses",
                "resume_insights",
            ])

            return Response(
                {
                    "message": "Resume uploaded successfully.",
                    "resume_id": resume_data.id,
                    "is_resume_uploaded": True,
                    "resume_details": extracted_data,
                    "recommended_jobs": [],
                    "search_query": "",
                    "provider_responses": {},
                    "resume_insights": resume_insights,
                },
                status=status.HTTP_200_OK
            )

        except Exception as e:

            resume_data.delete()

            return Response(
                {
                    "message": str(e)
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    def extract_pdf_text(self, path):
        """
        Extract text from PDF using PyMuPDF, preserving reading order.
        """
        text = ""
        pdf = fitz.open(path)

        for page in pdf:
            blocks = page.get_text("blocks")
            # Sort top-to-bottom, then left-to-right
            blocks.sort(key=lambda b: (round(b[1], 1), round(b[0], 1)))
            for b in blocks:
                block_text = b[4].strip()
                if block_text:
                    text += block_text + "\n"

        pdf.close()
        return text


    def extract_docx_text(self, path):
        """
        Extract text from DOCX, including tables and headers.
        """
        document = Document(path)
        text_parts = []

        # Headers often contain name/contact info
        for section in document.sections:
            header = section.header
            for para in header.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

        # Body paragraphs, in document order
        for element in document.element.body:
            if element.tag.endswith('p'):
                for para in document.paragraphs:
                    if para._p == element and para.text.strip():
                        text_parts.append(para.text.strip())
            elif element.tag.endswith('tbl'):
                for table in document.tables:
                    if table._tbl == element:
                        for row in table.rows:
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells if cell.text.strip()
                            )
                            if row_text:
                                text_parts.append(row_text)

        return "\n".join(text_parts)


    def extract_resume_details(self, text):
        """
        Extract structured resume information using Groq, with retries
        for transient failures or malformed JSON.
        """

        if not text or not text.strip():
            print("Resume text extraction returned empty text.")
            return self.normalize_resume_details({}, text)

        if client is None:
            return self.normalize_resume_details({}, text)

        # Guard against extremely long resumes blowing the context window,
        # which can cause truncation and inconsistent field extraction.
        max_chars = 15000
        truncated_text = text[:max_chars]

        prompt = f"""
Extract structured information from the resume below.

Rules:
- Return ONLY valid JSON.
- Do not add markdown or explanations.
- Do not hallucinate values.
- Use "" for missing strings.
- Use [] for missing arrays.
- Use 0 for missing numbers.

JSON Schema:
{{
    "name": "",
    "email": "",
    "phone": "",
    "skills": [],
    "experience": "",
    "experience_years": 0,
    "education": [],
    "projects": [],
    "certifications": [],
    "current_company": "",
    "current_designation": "",
    "location": "",
    "linkedin": "",
    "github": ""
}}

Resume:

{truncated_text}
"""

        last_error = None

        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are an ATS Resume Parser. Return only valid JSON, matching the schema exactly."
                        },
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    temperature=0,
                    response_format={"type": "json_object"},
                )
                result = response.choices[0].message.content.strip()

                if result.startswith("```"):
                    result = result.strip("`")
                    if result.startswith("json"):
                        result = result[4:]
                    result = result.strip()

                parsed = self.safe_json_loads(result)

                if not isinstance(parsed, dict):
                    raise ValueError(f"Non-dict result: {type(parsed)}")

                # Basic sanity check: reject obviously empty extractions
                # rather than silently accepting a bad parse as "success".
                if not any(parsed.get(k) for k in ("name", "email", "skills", "experience")):
                    raise ValueError("Parsed result has no meaningful fields")

                return self.normalize_resume_details(parsed, text)

            except Exception as e:
                last_error = e
                print(f"Groq Error (attempt {attempt + 1}/3):", str(e))
                if attempt < 2:
                    time.sleep(1.5 * (attempt + 1))

        print("Resume parsing failed after 3 attempts:", str(last_error))
        return self.normalize_resume_details({}, text)

    def _generate_resume_insights_with_openai(self, extracted_data, recommended_jobs=None):
        recommended_jobs = recommended_jobs or []
        skills = extracted_data.get("skills") or []
        if not isinstance(skills, list):
            skills = [skills]

        normalized_skills = [str(s).strip() for s in skills if isinstance(s, str) and s.strip()]
        if not normalized_skills:
            normalized_skills = ["Python", "Django", "React"]

        designation = extracted_data.get("current_designation") or "Software Developer"
        experience_years = extracted_data.get("experience_years", 0)

        # Pick top 4 skills most relevant to the designation (skip generic tools)
        skip_tools = {"git", "github", "vs code", "postman", "chrome devtools", "bitbucket", "jira", "slack", "figma"}
        core_skills = [s for s in normalized_skills if s.lower() not in skip_tools][:4]
        if not core_skills:
            core_skills = normalized_skills[:4]

        prompt = f"""You are a Senior Technical Interview Coach. Generate a complete interview preparation pack for this candidate.

Designation: {designation}
Skills: {json.dumps(core_skills)}
Experience: {experience_years} years

For EACH skill in the list above, generate EXACTLY:
- 3 Basic questions with full answers
- 3 Intermediate questions with full answers
- 3 Advanced questions with full answers
- 3 Coding questions with complete working code solutions

Total = 12 items per skill.

For Coding level: the question must be a real coding problem (e.g. implement a function, fix a bug, write a component). The answer must contain the COMPLETE working code solution with explanation.

For Basic/Intermediate/Advanced: the answer must be a thorough, interview-ready explanation — not a hint. Write as if explaining to an interviewer at Google/Amazon.

Return ONLY this JSON (no markdown, no explanation outside JSON):
{{
  "ats_resume_score": 82,
  "resume_improvement_suggestions": [
    "Add measurable impact metrics to each work experience bullet point.",
    "Include a concise professional summary tailored to {designation} roles.",
    "List certifications with issuing body and year obtained.",
    "Quantify project outcomes (e.g. reduced load time by 40%).",
    "Add relevant keywords from job descriptions to pass ATS filters."
  ],
  "ai_interview_preparation": [
    {{"skill": "React.js", "level": "Basic", "question": "What is the virtual DOM in React and how does it work?", "answer": "The virtual DOM is a lightweight in-memory representation of the real DOM. When state changes, React creates a new virtual DOM tree, diffs it against the previous one (reconciliation), and only updates the changed nodes in the real DOM. This avoids expensive full-page re-renders. Example: when you call setState, React re-renders the component in the virtual DOM first, computes the minimal set of changes, then applies only those to the browser DOM."}},
    {{"skill": "React.js", "level": "Coding", "question": "Write a custom React hook useDebounce(value, delay) that delays updating the returned value until the user stops typing.", "answer": "import {{ useState, useEffect }} from 'react';\n\nfunction useDebounce(value, delay) {{\n  const [debouncedValue, setDebouncedValue] = useState(value);\n\n  useEffect(() => {{\n    const timer = setTimeout(() => {{\n      setDebouncedValue(value);\n    }}, delay);\n\n    return () => clearTimeout(timer); // cleanup on value change\n  }}, [value, delay]);\n\n  return debouncedValue;\n}}\n\n// Usage:\n// const debouncedSearch = useDebounce(searchTerm, 500);"}}
  ]
}}

Now generate the full response for ALL skills: {json.dumps(core_skills)}. Every skill must have exactly 3 Basic + 3 Intermediate + 3 Advanced + 3 Coding items."""

        skills_json = json.dumps(core_skills)
        prompt = f"""Create an interview prep pack from this resume data.

Designation: {designation}
Skills: {skills_json}
Experience: {experience_years} years

Rules:
- Return only valid JSON.
- For each skill, generate exactly 3 Basic, 3 Intermediate, 3 Advanced, and 3 Coding items.
- Each item must have: skill, level, question, answer.
- Coding answers must include complete working code plus a brief explanation.
- Non-coding answers must be clear, accurate, and interview-ready.
- Use only the skills above; do not invent extra skills.
- Return exactly 5 short, actionable resume improvement suggestions.
- No markdown, no extra text.

JSON shape:
{{
  "ats_resume_score": 82,
  "resume_improvement_suggestions": ["..."],
  "ai_interview_preparation": [{{"skill": "", "level": "", "question": "", "answer": ""}}]
}}

Generate items for all skills in {skills_json}."""

        try:
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a Senior Technical Interview Coach. Return only valid JSON with no markdown or extra text.",
                    },
                    {
                        "role": "user",
                        "content": prompt,
                    },
                ],
                temperature=0.2,
                response_format={"type": "json_object"},
                max_tokens=8000,
            )

            content = response.choices[0].message.content.strip()
            parsed = self.safe_json_loads(content)

            if not isinstance(parsed, dict):
                raise ValueError("Groq insight response was not a JSON object.")

            ats_score = parsed.get("ats_resume_score")
            suggestions = parsed.get("resume_improvement_suggestions")
            interview_prep = parsed.get("ai_interview_preparation")

            if not isinstance(ats_score, int):
                try:
                    ats_score = int(float(ats_score))
                except (TypeError, ValueError):
                    ats_score = 0

            if not isinstance(suggestions, list) or len(suggestions) < 3:
                raise ValueError("Groq suggestions payload is incomplete.")

            if not isinstance(interview_prep, list) or len(interview_prep) < 5:
                raise ValueError("Groq interview prep payload is incomplete.")

            flat_prep = []
            for item in interview_prep:
                if not isinstance(item, dict):
                    continue
                answer = str(item.get("answer") or item.get("tip") or "").strip()
                if all(k in item for k in ("skill", "level", "question")) and answer:
                    flat_prep.append({
                        "skill": str(item["skill"]).strip(),
                        "level": str(item["level"]).strip(),
                        "question": str(item["question"]).strip(),
                        "answer": answer,
                    })

            if len(flat_prep) < 5:
                raise ValueError("Groq interview prep items are not in the expected flat format.")

            return {
                "ats_resume_score": max(0, min(100, ats_score)),
                "resume_improvement_suggestions": suggestions[:5],
                "ai_interview_preparation": flat_prep,
            }
        except Exception as exc:
            print("Groq resume insight generation failed:", str(exc))
            return None

    def generate_resume_insights(self, extracted_data, recommended_jobs=None):
        recommended_jobs = recommended_jobs or []
        openai_insights = self._generate_resume_insights_with_openai(extracted_data, recommended_jobs)
        if isinstance(openai_insights, dict):
            return openai_insights

        skills = extracted_data.get("skills") or []
        if not isinstance(skills, list):
            skills = [skills]

        normalized_skills = [str(skill).strip() for skill in skills if isinstance(skill, str) and skill.strip()]
        if not normalized_skills:
            normalized_skills = ["python", "django", "react"]

        matched_skill_count = 0
        for job in recommended_jobs:
            if not isinstance(job, dict):
                continue
            job_text = " ".join([
                str(job.get("title") or ""),
                str(job.get("description") or ""),
                str(job.get("company") or ""),
                str(job.get("matched_skills") or ""),
            ]).lower()
            for skill in normalized_skills:
                if skill.lower() in job_text:
                    matched_skill_count += 1
                    break

        job_match_ratio = (matched_skill_count / max(len(normalized_skills), 1)) * 100 if recommended_jobs else 0
        ats_score = int(round(min(98, max(55, 68 + (job_match_ratio * 0.25)))))

        missing_skills = []
        for skill in normalized_skills:
            if not any(skill.lower() in str(job.get("title", "")).lower() or skill.lower() in str(job.get("description", "")).lower() for job in recommended_jobs if isinstance(job, dict)):
                missing_skills.append(skill)

        resume_improvement_suggestions = [
            f"Add quantified proof for {skill} in your resume bullet points." for skill in missing_skills[:3]
        ]
        if not resume_improvement_suggestions:
            resume_improvement_suggestions = [
                "Keep the resume outcome-focused by adding measurable project metrics and business impact.",
                "Highlight leadership, collaboration, and delivery ownership in a concise format.",
                "Tailor one resume version to the target role with relevant keywords and tools.",
            ]

        interview_prep = []
        for skill in normalized_skills[:4]:
            s = skill
            for level, question, answer in [
                ("Basic", f"What is {s} and what problem does it solve?",
                 f"{s} is a widely-used technology in modern software development. It solves problems like code reusability, performance, and developer productivity. For example, React solves the problem of efficiently updating the UI by using a virtual DOM. Django solves rapid backend development by providing an ORM, admin panel, and authentication out of the box."),
                ("Basic", f"What are the core concepts you must know in {s}?",
                 f"Core concepts in {s}: 1) Component/module structure and how they communicate, 2) State and data management, 3) Error handling patterns, 4) Configuration and setup, 5) Testing approach. Mastering these gives you a strong foundation for any interview question on {s}."),
                ("Basic", f"How do you install and set up {s} in a new project?",
                 f"For Python packages: pip install {s}. For JS: npm install {s}. After installing, import it at the top of your file and follow the official quickstart. Always pin the version in requirements.txt or package.json to avoid breaking changes."),
                ("Intermediate", f"How does {s} handle performance optimization?",
                 f"{s} provides several built-in performance tools. For React: use React.memo to prevent unnecessary re-renders, useMemo/useCallback to memoize values and functions, and lazy() with Suspense for code splitting. For Django: use select_related() and prefetch_related() to reduce database queries, and cache views with Django's cache framework."),
                ("Intermediate", f"How do you handle errors and exceptions in {s}?",
                 f"In Python/{s}: wrap risky operations in try/except blocks and raise specific exceptions. In JavaScript/{s}: use try/catch with async/await, and add error boundaries in React components. Always log errors with context and return user-friendly messages. Never expose stack traces to end users."),
                ("Intermediate", f"Explain a real-world architecture pattern you used with {s}.",
                 f"I used {s} in a layered architecture: the presentation layer handled UI/API responses, the service layer contained business logic, and the data layer managed database access. This separation made the codebase testable and maintainable. For example, in a Django REST API, views call service functions which call repository functions — no business logic in views."),
                ("Advanced", f"How would you scale a {s} application to handle 1 million users?",
                 f"Scaling {s} to 1M users: 1) Horizontal scaling — run multiple instances behind a load balancer, 2) Database optimization — read replicas, connection pooling, query caching, 3) CDN for static assets, 4) Async task queues (Celery/Redis) for heavy operations, 5) Rate limiting and circuit breakers to protect services. Monitor with Prometheus + Grafana."),
                ("Advanced", f"How do you secure a {s} application against common vulnerabilities?",
                 f"Security checklist for {s}: 1) SQL injection — use ORM/parameterized queries, never raw string interpolation, 2) XSS — sanitize and escape all user input, 3) CSRF — use CSRF tokens on state-changing requests, 4) Auth — use JWT with short expiry + refresh tokens, 5) Secrets — never hardcode API keys, use environment variables. Run OWASP ZAP scans regularly."),
                ("Advanced", f"Describe a production incident you resolved involving {s}.",
                 f"Situation: Our {s} service started returning 500 errors under load. Task: Identify root cause within 30 minutes. Action: Checked logs — found N+1 query problem causing DB connection pool exhaustion. Added select_related() to the queryset and deployed a hotfix. Result: Response time dropped from 8s to 200ms, zero errors for 60+ days. Lesson: always profile queries before going to production."),
                ("Coding", f"Write a function in {s} to find the two numbers in an array that sum to a target value.",
                 f"# Python solution using hash map — O(n) time, O(n) space\ndef two_sum(nums, target):\n    seen = {{}}\n    for i, num in enumerate(nums):\n        complement = target - num\n        if complement in seen:\n            return [seen[complement], i]\n        seen[num] = i\n    return []\n\n# Example:\nprint(two_sum([2, 7, 11, 15], 9))  # Output: [0, 1]\n\n# JavaScript equivalent:\nfunction twoSum(nums, target) {{\n  const seen = {{}};\n  for (let i = 0; i < nums.length; i++) {{\n    const complement = target - nums[i];\n    if (complement in seen) return [seen[complement], i];\n    seen[nums[i]] = i;\n  }}\n  return [];\n}}"),
                ("Coding", f"Implement a debounce function for {s} that delays execution until the user stops triggering it.",
                 f"// JavaScript debounce — used in search inputs, resize handlers\nfunction debounce(fn, delay) {{\n  let timer;\n  return function(...args) {{\n    clearTimeout(timer);\n    timer = setTimeout(() => fn.apply(this, args), delay);\n  }};\n}}\n\n// React hook version:\nimport {{ useState, useEffect }} from 'react';\nfunction useDebounce(value, delay) {{\n  const [debounced, setDebounced] = useState(value);\n  useEffect(() => {{\n    const t = setTimeout(() => setDebounced(value), delay);\n    return () => clearTimeout(t);\n  }}, [value, delay]);\n  return debounced;\n}}\n\n// Python equivalent using threading:\nimport threading\ndef debounce(fn, delay):\n    timer = None\n    def wrapper(*args, **kwargs):\n        nonlocal timer\n        if timer: timer.cancel()\n        timer = threading.Timer(delay, fn, args, kwargs)\n        timer.start()\n    return wrapper"),
                ("Coding", f"Write a {s} function to flatten a deeply nested array/list into a single flat list.",
                 f"# Python — recursive approach\ndef flatten(arr):\n    result = []\n    for item in arr:\n        if isinstance(item, list):\n            result.extend(flatten(item))\n        else:\n            result.append(item)\n    return result\n\nprint(flatten([1, [2, [3, [4]], 5]]))  # [1, 2, 3, 4, 5]\n\n# Python one-liner using itertools:\nfrom itertools import chain\ndef flatten_iter(arr):\n    return list(chain.from_iterable(x if isinstance(x, list) else [x] for x in arr))\n\n// JavaScript:\nconst flatten = (arr) => arr.reduce((acc, val) =>\n  Array.isArray(val) ? acc.concat(flatten(val)) : acc.concat(val), []);\n\nconsole.log(flatten([1, [2, [3, [4]], 5]])); // [1, 2, 3, 4, 5]"),
            ]:
                interview_prep.append({"skill": s, "level": level, "question": question, "answer": answer})

        return {
            "ats_resume_score": ats_score,
            "resume_improvement_suggestions": resume_improvement_suggestions,
            "ai_interview_preparation": interview_prep,
        }

    def safe_json_loads(self, raw_text):
        if not isinstance(raw_text, str):
            return raw_text

        cleaned = raw_text.strip()
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            cleaned = cleaned[start : end + 1]

        return json.loads(cleaned)

    def extract_first_match(self, pattern, text, flags=re.IGNORECASE):
        match = re.search(pattern, text, flags)
        if not match:
            return ""
        value = match.group(1).strip()
        return value.rstrip(".,;")

    def extract_section_items(self, lines, section_name):
        section_headers = {
            "experience",
            "skills",
            "projects",
            "education",
            "certifications",
            "summary",
            "objective",
            "profile",
            "work experience",
            "professional experience",
            "technical skills",
            "contact",
            "academic background",
            "academic qualifications",
            "education details",
            "project experience",
            "personal projects",
            "technical projects",
            "certificates",
            "awards",
        }

        section_index = -1
        target = section_name.lower()
        for index, line in enumerate(lines):
            normalized = line.lower().rstrip(":")
            if normalized == target or normalized.startswith(f"{target} "):
                section_index = index
                break
            if target in {"education", "projects", "certifications"} and normalized in {
                "academic background",
                "academic qualifications",
                "education details",
                "project experience",
                "personal projects",
                "technical projects",
                "certificates",
            }:
                if normalized.startswith(target) or (target == "education" and normalized in {"academic background", "academic qualifications", "education details"}) or (target == "projects" and normalized in {"project experience", "personal projects", "technical projects"}) or (target == "certifications" and normalized in {"certificates"}):
                    section_index = index
                    break

        if section_index == -1:
            return []

        items = []
        for line in lines[section_index + 1 :]:
            normalized = line.lower().rstrip(":")
            if normalized in section_headers:
                break
            if re.match(r"^[•\-\*]\s*", line):
                cleaned = re.sub(r"^[•\-\*]\s*", "", line).strip()
                if cleaned:
                    items.append(cleaned)
            elif line and not re.match(r"^(name|email|phone|address|location|linkedin|github):", line, re.IGNORECASE):
                items.append(line)

        deduped = []
        for item in items:
            if item and item.lower() not in {existing.lower() for existing in deduped}:
                deduped.append(item)
        return deduped[:8]

    def fallback_resume_details(self, text):
        lines = [line.strip(" \t-•*") for line in text.splitlines()]
        lines = [line for line in lines if line]
        joined = "\n".join(lines)
        lower_joined = joined.lower()

        designation_aliases = [
            ("front-end developer", "Front-End Developer"),
            ("frontend developer", "Front-End Developer"),
            ("react js developer", "React JS Developer"),
            ("react developer", "React Developer"),
            ("full stack developer", "Full Stack Developer"),
            ("software engineer", "Software Engineer"),
            ("software developer", "Software Developer"),
            ("python developer", "Python Developer"),
            ("backend developer", "Backend Developer"),
            ("node js developer", "Node.js Developer"),
            ("java developer", "Java Developer"),
            ("ui developer", "UI Developer"),
            ("ui/ux designer", "UI/UX Designer"),
            ("devops engineer", "DevOps Engineer"),
            ("data analyst", "Data Analyst"),
            ("data scientist", "Data Scientist"),
            ("product manager", "Product Manager"),
        ]

        skill_aliases = [
            ("react.js", "React.js"),
            ("reactjs", "React.js"),
            ("react", "React"),
            ("javascript", "JavaScript"),
            ("typescript", "TypeScript"),
            ("next.js", "Next.js"),
            ("nextjs", "Next.js"),
            ("node.js", "Node.js"),
            ("nodejs", "Node.js"),
            ("python", "Python"),
            ("django rest framework", "Django REST Framework"),
            ("django", "Django"),
            ("rest api", "REST API"),
            ("bootstrap", "Bootstrap"),
            ("material ui", "Material UI"),
            ("mui", "MUI"),
            ("html5", "HTML5"),
            ("css3", "CSS3"),
            ("git", "Git"),
            ("github", "GitHub"),
            ("bitbucket", "Bitbucket"),
            ("postman", "Postman"),
            ("agile", "Agile"),
            ("scrum", "Scrum"),
            ("crud", "CRUD"),
            ("form validation", "Form Validation"),
            ("lazy loading", "Lazy Loading"),
            ("performance optimization", "Performance Optimization"),
            ("cross-browser compatibility", "Cross-browser Compatibility"),
            ("component-based architecture", "Component-Based Architecture"),
        ]

        email = self.extract_first_match(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})", joined)
        phone = self.extract_first_match(r"(\+?\d[\d\s().-]{8,}\d)", joined)
        linkedin = self.extract_first_match(r"(https?://(?:www\.)?linkedin\.com/[^\s)]+)", joined)
        github = self.extract_first_match(r"(https?://(?:www\.)?github\.com/[^\s)]+)", joined)

        designation = ""
        for line in lines[:25]:
            lowered = line.lower()
            if any(alias in lowered for alias, _ in designation_aliases):
                for alias, label in designation_aliases:
                    if alias in lowered:
                        designation = label
                        break
                if designation:
                    break

        if not designation and lines:
            designation = lines[0]

        location = self.extract_first_match(r"(?:location|address)\s*[:\-]\s*([^\n]+)", joined)
        if not location:
            location_candidates = [
                line
                for line in lines[:20]
                if "," in line
                and len(line.split()) <= 5
                and not re.search(r"@|linkedin|github|\d{10,}", line, re.IGNORECASE)
            ]
            location = location_candidates[0] if location_candidates else ""

        experience_years = 0
        year_matches = re.findall(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)", lower_joined, re.IGNORECASE)
        if year_matches:
            try:
                experience_years = int(float(max(year_matches, key=lambda item: float(item))))
            except ValueError:
                experience_years = 0

        experience = self.extract_first_match(r"(?:experience|work experience)[:\s-]*([^\n]+)", joined)
        if not experience and experience_years:
            experience = f"{experience_years} years"

        skills = []
        for alias, label in skill_aliases:
            if alias in lower_joined:
                skills.append(label)
        normalized_skills = []
        for skill in skills:
            if skill.lower() not in {item.lower() for item in normalized_skills}:
                normalized_skills.append(skill)

        education = self.extract_section_items(lines, "education")
        projects = self.extract_section_items(lines, "projects")
        certifications = self.extract_section_items(lines, "certifications")

        return {
            "name": lines[0] if lines else "",
            "email": email,
            "phone": phone,
            "skills": normalized_skills,
            "experience": experience,
            "experience_years": experience_years,
            "education": education,
            "projects": projects,
            "certifications": certifications,
            "current_company": "",
            "current_designation": designation,
            "location": location,
            "linkedin": linkedin,
            "github": github,
        }

    def normalize_resume_details(self, parsed_data, text):
        fallback = self.fallback_resume_details(text or "")
        if not isinstance(parsed_data, dict):
            return fallback

        normalized = dict(fallback)
        for key, value in parsed_data.items():
            if key in {"skills", "education", "projects", "certifications"}:
                if isinstance(value, list) and value:
                    normalized[key] = [item for item in value if item not in (None, "", [])]
            elif key == "experience_years":
                try:
                    normalized[key] = int(float(value))
                except (TypeError, ValueError):
                    pass
            elif isinstance(value, str) and value.strip():
                normalized[key] = value.strip()
            elif value not in (None, "", []):
                normalized[key] = value

        if not normalized.get("current_designation"):
            normalized["current_designation"] = (
                parsed_data.get("current_designation")
                if isinstance(parsed_data.get("current_designation"), str) and parsed_data.get("current_designation").strip()
                else parsed_data.get("designation")
                if isinstance(parsed_data.get("designation"), str) and parsed_data.get("designation").strip()
                else parsed_data.get("title")
                if isinstance(parsed_data.get("title"), str) and parsed_data.get("title").strip()
                else parsed_data.get("role")
                if isinstance(parsed_data.get("role"), str) and parsed_data.get("role").strip()
                else fallback.get("current_designation", "")
            )

        if not normalized.get("experience") and normalized.get("experience_years"):
            normalized["experience"] = f'{normalized["experience_years"]} years'

        return normalized
            
    
    def normalize_skills(self, skills):
        if not isinstance(skills, list):
            return []

        normalized = []
        for skill in skills:
            if not isinstance(skill, str):
                continue
            cleaned = skill.strip()
            if cleaned and cleaned.lower() not in {item.lower() for item in normalized}:
                normalized.append(cleaned)
        return normalized

    def build_job_query(self, extracted_data, filters=None):
        filters = filters or {}

        custom_query = filters.get("query") or filters.get("title")
        if isinstance(custom_query, str) and custom_query.strip():
            query = custom_query.strip()
        else:
            designation = (
                extracted_data.get("current_designation") or
                extracted_data.get("designation") or
                extracted_data.get("title") or
                extracted_data.get("role")
            )
            if isinstance(designation, str) and designation.strip():
                query = designation.strip()
            else:
                query = "Software Developer"

        location = filters.get("location")
        if isinstance(location, str) and location.strip() and location.strip().lower() != "all":
            query = f"{query} in {location.strip()}"

        return query

    def build_api_params(self, extracted_data, filters=None):
        filters = filters or {}
        num_pages = filters.get("num_pages", 20)
        try:
            num_pages = max(1, min(int(num_pages), 20))
        except (TypeError, ValueError):
            num_pages = 20

        work_from_home = filters.get("work_from_home")
        if isinstance(work_from_home, str):
            work_from_home = work_from_home.strip().lower() in {"true", "1", "yes", "on"}

        params = {
            "query": self.build_job_query(extracted_data, filters or {}),
            "country": "in",
            "num_pages": num_pages,
        }

        if filters.get("date_posted") in {"today", "3days", "week", "month"}:
            params["date_posted"] = filters["date_posted"]
        elif filters.get("posted") == "7":
            params["date_posted"] = "week"
        elif filters.get("posted") == "30":
            params["date_posted"] = "month"

        if isinstance(work_from_home, bool):
            params["work_from_home"] = work_from_home

        employment_type = filters.get("employment_types") or filters.get("employment_type")
        if isinstance(employment_type, str):
            normalized = employment_type.strip().lower()
            employment_map = {
                "full-time": "FULLTIME",
                "fulltime": "FULLTIME",
                "contract": "CONTRACTOR",
                "contractor": "CONTRACTOR",
                "part-time": "PARTTIME",
                "parttime": "PARTTIME",
                "intern": "INTERN",
                "internship": "INTERN",
            }
            mapped = employment_map.get(normalized)
            if mapped:
                params["employment_types"] = mapped

        job_requirements = filters.get("job_requirements")
        if isinstance(job_requirements, str) and job_requirements.strip():
            params["job_requirements"] = job_requirements.strip()
        else:
            experience_filter = filters.get("experience")
            if experience_filter == "0-2":
                params["job_requirements"] = "under_3_years_experience"
            elif experience_filter == "3-5":
                params["job_requirements"] = "more_than_3_years_experience"
            elif experience_filter == "5+":
                params["job_requirements"] = "more_than_3_years_experience"

        return params

    def _safe_get_json(self, url, headers=None, params=None, timeout=15):
        try:
            response = requests.get(
                url,
                headers=headers or {},
                params=params or {},
                timeout=timeout,
            )
            try:
                payload = response.json()
            except ValueError:
                payload = {}

            return {
                "status_code": response.status_code,
                "ok": response.ok,
                "payload": payload,
                "raw_text": response.text,
            }
        except Exception as exc:
            return {
                "status_code": None,
                "ok": False,
                "payload": {},
                "raw_text": str(exc),
            }

    def _collect_jobs_from_payload(self, payload):
        jobs = []

        if isinstance(payload, dict):
            data = payload.get("data")
            if isinstance(data, dict):
                nested_jobs = data.get("jobs")
                if isinstance(nested_jobs, list):
                    jobs.extend([item for item in nested_jobs if isinstance(item, dict)])
                elif isinstance(data.get("results"), list):
                    jobs.extend([item for item in data.get("results", []) if isinstance(item, dict)])
                elif isinstance(data, dict):
                    jobs.append(data)
            elif isinstance(payload.get("jobs"), list):
                jobs.extend([item for item in payload.get("jobs", []) if isinstance(item, dict)])
            elif isinstance(payload.get("results"), list):
                jobs.extend([item for item in payload.get("results", []) if isinstance(item, dict)])
            elif isinstance(payload.get("job_results"), list):
                jobs.extend([item for item in payload.get("job_results", []) if isinstance(item, dict)])
            elif isinstance(payload.get("organic_results"), list):
                jobs.extend([item for item in payload.get("organic_results", []) if isinstance(item, dict)])
            elif isinstance(payload.get("jobs_results"), list):
                jobs.extend([item for item in payload.get("jobs_results", []) if isinstance(item, dict)])
        elif isinstance(payload, list):
            jobs.extend([item for item in payload if isinstance(item, dict)])

        return jobs

    def _normalize_recommended_job(self, job, resume_skills):
        if not isinstance(job, dict):
            return None

        title = job.get("job_title") or job.get("title") or ""
        description = job.get("job_description") or job.get("description") or ""
        apply_link = (
            job.get("job_apply_link")
            or job.get("apply_link")
            or job.get("job_apply_url")
            or job.get("url")
            or job.get("link")
        )
        if not apply_link and isinstance(job.get("apply_options"), list):
            for option in job.get("apply_options", []):
                if isinstance(option, dict):
                    apply_link = option.get("apply_link") or option.get("url")
                    if apply_link:
                        break

        job_text = (title + " " + description).lower()
        matched_skills = [skill for skill in resume_skills if skill in job_text]

        skill_score = 0
        if resume_skills:
            skill_score = (len(matched_skills) / len(resume_skills)) * 100

        return {
            **job,
            "title": title,
            "company": job.get("employer_name") or job.get("company") or job.get("company_name") or "",
            "location": job.get("job_location") or job.get("job_city") or job.get("location") or "",
            "employment_type": job.get("job_employment_type") or job.get("job_employment_types") or job.get("employment_type") or "",
            "work_mode": job.get("job_work_from_home") or job.get("job_work_mode") or job.get("work_mode") or "",
            "experience": job.get("job_required_experience") or job.get("job_experience") or job.get("experience") or "",
            "salary": job.get("job_salary_string") or job.get("job_salary") or job.get("salary") or "",
            "posted_at": job.get("job_posted_at") or job.get("posted_at") or job.get("job_posted_date") or "",
            "description": description,
            "apply_link": apply_link,
            "url": apply_link,
            "matched_skills": matched_skills,
            "match_score": round(skill_score, 2),
        }

    def search_jobs(self, extracted_data, filters=None):

        params = self.build_api_params(
            extracted_data,
            filters,
        )
        query = params["query"]
        skills = extracted_data.get("skills", [])
        resume_skills = {
            skill.lower().replace(".js", "").strip()
            for skill in skills
            if isinstance(skill, str)
        }

        provider_responses = {}
        all_recommended_jobs = []

        rapidapi_host = os.getenv("RAPIDAPI_HOST")
        rapidapi_key = os.getenv("RAPIDAPI_KEY") 
        rapidapi_url = "https://jsearch.p.rapidapi.com/search-v2"
        rapidapi_headers = {
            "Content-Type": "application/json",
            "x-rapidapi-host": rapidapi_host,
            "x-rapidapi-key": rapidapi_key,
        }

        rapidapi_params = {
           "time_frame": "7d", "query": query ,"num_pages" :"20", "country":"in", "date_posted":"month"
        }

        rapidapi_response = self._safe_get_json(
            rapidapi_url,
            headers=rapidapi_headers,
            params=rapidapi_params,
        )
        provider_responses["rapidapi_jsearch"] = {
            "status_code": rapidapi_response.get("status_code"),
            "ok": rapidapi_response.get("ok"),
            "raw_response": rapidapi_response.get("payload"),
        }

        if rapidapi_response.get("ok"):
            jobs = self._collect_jobs_from_payload(rapidapi_response.get("payload"))
            for job in jobs:
                normalized_job = self._normalize_recommended_job(job, resume_skills)
                if normalized_job:
                    all_recommended_jobs.append(normalized_job)

        fantastic_token = os.getenv("FANTASTIC_JOBS_TOKEN")
        fantastic_response = self._safe_get_json(
            "https://data.fantastic.jobs/v1/active-ats",
            headers={"Authorization": f"Bearer {fantastic_token}"} if fantastic_token else {},
            params={"time_frame": "7d", "title": query ,"limit" :"500"},
        )
        provider_responses["fantastic_jobs"] = {
            "status_code": fantastic_response.get("status_code"),
            "ok": fantastic_response.get("ok"),
            "raw_response": fantastic_response.get("payload"),
        }
        if fantastic_response.get("ok"):
            for job in self._collect_jobs_from_payload(fantastic_response.get("payload")):
                normalized_job = self._normalize_recommended_job(job, resume_skills)
                if normalized_job:
                    all_recommended_jobs.append(normalized_job)

        serpapi_key = os.getenv("SERPAPI_API_KEY")
        serpapi_params = {
            "engine": "google_jobs",
            "location": "India",
            "google_domain": "google.co.in",
            "hl": "hi",
            "gl": "in",
            "q": query,
            "num": 10,
        }

        serpapi_response = {
            "status_code": None,
            "ok": False,
            "payload": {},
            "raw_text": "Missing SERPAPI_API_KEY",
        }

        if serpapi_key:
            max_pages = 5
            serpapi_jobs = []
            next_page_token = None
            page_index = 1

            while page_index <= max_pages:
                current_params = dict(serpapi_params)
                current_params["api_key"] = serpapi_key
                if next_page_token:
                    current_params["next_page_token"] = next_page_token

                current_response = self._safe_get_json(
                    "https://serpapi.com/search",
                    params=current_params,
                )

                if not current_response.get("ok"):
                    serpapi_response = current_response
                    break

                current_payload = current_response.get("payload") or {}
                page_jobs = self._collect_jobs_from_payload(current_payload)
                if not page_jobs:
                    break

                serpapi_jobs.extend(page_jobs)
                serpapi_response = current_response

                pagination = current_payload.get("serpapi_pagination") if isinstance(current_payload, dict) else {}
                next_page_token = pagination.get("next_page_token") if isinstance(pagination, dict) else None
                if not next_page_token:
                    break

                page_index += 1

            serpapi_response["payload"] = {
                "jobs_results": serpapi_jobs,
                "search_metadata": serpapi_response.get("payload", {}).get("search_metadata", {}),
                "serpapi_pagination": serpapi_response.get("payload", {}).get("serpapi_pagination", {}),
            }
            serpapi_response["raw_text"] = json.dumps(serpapi_response.get("payload"))

        provider_responses["serpapi"] = {
            "status_code": serpapi_response.get("status_code"),
            "ok": serpapi_response.get("ok"),
            "raw_response": serpapi_response.get("payload"),
        }
        if serpapi_response.get("ok"):
            for job in self._collect_jobs_from_payload(serpapi_response.get("payload")):
                normalized_job = self._normalize_recommended_job(job, resume_skills)
                if normalized_job:
                    all_recommended_jobs.append(normalized_job)

        deduped_jobs = []
        seen_keys = set()
        for job in all_recommended_jobs:
            job_key = (
                job.get("apply_link")
                or job.get("url")
                or job.get("job_apply_link")
                or f"{job.get('title')}::{job.get('company')}::{job.get('location')}"
            )
            if not job_key:
                continue
            if job_key in seen_keys:
                continue
            seen_keys.add(job_key)
            deduped_jobs.append(job)

        deduped_jobs.sort(
            key=lambda x: x["match_score"],
            reverse=True,
        )

        return deduped_jobs, query, provider_responses


class JobSearchAPIView(APIView):

    def get(self, request):
        user_id = request.query_params.get("user_id")
        if not user_id:
            return Response(
                {"message": "user_id is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response(
                {"message": "User not found."},
                status=status.HTTP_404_NOT_FOUND,
            )

        persisted_resume = Resume.objects.filter(user=user).order_by("-updated_at", "-created_at").first()
        if not persisted_resume:
            return Response(
                {
                    "is_resume_uploaded": False,
                    "resume_details": {},
                    "recommended_jobs": [],
                    "resume_insights": {},
                    "query": "",
                    "provider_responses": {},
                },
                status=status.HTTP_200_OK,
            )

        return Response(
            {
                "is_resume_uploaded": bool(persisted_resume.is_resume_uploaded),
                "resume_details": persisted_resume.resume_details or {},
                "recommended_jobs": persisted_resume.last_recommended_jobs or [],
                "resume_insights": persisted_resume.resume_insights or {},
                "query": persisted_resume.search_query or "",
                "provider_responses": persisted_resume.provider_responses or {},
            },
            status=status.HTTP_200_OK,
        )

    def post(self, request):
        user_id = request.data.get("user_id")
        extracted_data = request.data.get("resume_details") or {}
        filters = request.data.get("filters") or {}
        persisted_resume = None

        if user_id:
            try:
                user = userModel.objects.get(id=user_id)
            except userModel.DoesNotExist:
                return Response(
                    {"message": "User not found."},
                    status=status.HTTP_404_NOT_FOUND,
                )

            persisted_resume = Resume.objects.filter(user=user).order_by("-updated_at", "-created_at").first()
            if persisted_resume and not extracted_data:
                extracted_data = persisted_resume.resume_details or {}
            if persisted_resume and not filters:
                filters = persisted_resume.last_search_filters or {}

        if not isinstance(extracted_data, dict):
            return Response(
                {"message": "resume_details must be an object."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        helper = ResumeUploadAPIView()
        service = ResumeAIService()

        try:
            recommended_jobs, query, provider_responses = helper.search_jobs(extracted_data, filters)
            resume_insights = service.generate_ats_score(extracted_data) or {}

            if user_id:
                if persisted_resume:
                    persisted_resume.resume_details = extracted_data
                    persisted_resume.last_search_filters = filters
                    persisted_resume.last_recommended_jobs = recommended_jobs
                    persisted_resume.search_query = query
                    persisted_resume.provider_responses = provider_responses
                    persisted_resume.resume_insights = resume_insights
                    persisted_resume.save(update_fields=[
                        "resume_details",
                        "last_search_filters",
                        "last_recommended_jobs",
                        "search_query",
                        "provider_responses",
                        "resume_insights",
                    ])

            return Response(
                {
                    "query": query,
                    "recommended_jobs": recommended_jobs,
                    "provider_responses": provider_responses,
                    "resume_details": extracted_data,
                    "resume_insights": resume_insights,
                },
                status=status.HTTP_200_OK,
            )
        except Exception as e:
            return Response(
                {"message": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class AppliedJobAPIView(APIView):

    def get(self, request):
        user_id = request.query_params.get("user_id")
        if not user_id:
            return Response(
                {"message": "user_id is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response(
                {"message": "User not found."},
                status=status.HTTP_404_NOT_FOUND,
            )

        applied_jobs = AppliedJob.objects.filter(user=user).order_by("-applied_at")
        serializer = AppliedJobSerializer(applied_jobs, many=True)
        return Response({"results": serializer.data}, status=status.HTTP_200_OK)

    def post(self, request):
        user_id = request.data.get("user_id")
        job_data = request.data.get("job") or request.data

        if not user_id:
            return Response(
                {"message": "user_id is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response(
                {"message": "User not found."},
                status=status.HTTP_404_NOT_FOUND,
            )

        if not isinstance(job_data, dict):
            return Response(
                {"message": "job data is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        apply_link = job_data.get("apply_link") or job_data.get("url") or ""
        title = job_data.get("title") or job_data.get("job_title") or ""
        company = job_data.get("company") or job_data.get("employer_name") or ""
        location = job_data.get("location") or job_data.get("job_location") or ""

        defaults = {
            "title": title,
            "company": company,
            "location": location,
            "employment_type": job_data.get("employment_type") or "",
            "work_mode": job_data.get("work_mode") or "",
            "experience": job_data.get("experience") or "",
            "posted_at": job_data.get("posted_at") or "",
            "summary": job_data.get("summary") or job_data.get("description") or "",
            "description": job_data.get("description") or "",
            "apply_link": apply_link,
            "match_score": float(job_data.get("match_score") or 0),
            "raw_data": job_data,
        }

        lookup = {"user": user}
        if apply_link:
            lookup["apply_link"] = apply_link
        elif title and company:
            lookup["title"] = title
            lookup["company"] = company

        applied_job, created = AppliedJob.objects.get_or_create(defaults=defaults, **lookup)
        if not created:
            serializer = AppliedJobSerializer(applied_job)
            return Response(
                {
                    "message": "Job already saved.",
                    "data": serializer.data,
                },
                status=status.HTTP_200_OK,
            )

        serializer = AppliedJobSerializer(applied_job)
        return Response(
            {
                "message": "Job applied successfully.",
                "data": serializer.data,
            },
            status=status.HTTP_201_CREATED,
        )


class SavedJobAPIView(APIView):

    def get(self, request):
        user_id = request.query_params.get("user_id")
        if not user_id:
            return Response({"message": "user_id is required."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response({"message": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        saved_jobs = SavedJob.objects.filter(user=user).order_by("-saved_at")
        serializer = SavedJobSerializer(saved_jobs, many=True)
        return Response({"results": serializer.data}, status=status.HTTP_200_OK)

    def post(self, request):
        user_id = request.data.get("user_id")
        job_data = request.data.get("job") or {}

        if not user_id:
            return Response({"message": "user_id is required."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response({"message": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        if not isinstance(job_data, dict):
            return Response({"message": "job data is required."}, status=status.HTTP_400_BAD_REQUEST)

        apply_link = job_data.get("apply_link") or job_data.get("url") or ""
        title = job_data.get("title") or ""
        company = job_data.get("company") or ""

        defaults = {
            "title": title,
            "company": company,
            "location": job_data.get("location") or "",
            "employment_type": job_data.get("employment_type") or "",
            "work_mode": job_data.get("work_mode") or "",
            "experience": job_data.get("experience") or "",
            "posted_at": job_data.get("posted_at") or "",
            "description": job_data.get("description") or "",
            "apply_link": apply_link,
            "match_score": float(job_data.get("match_score") or 0),
            "raw_data": job_data,
        }

        lookup = {"user": user}
        if apply_link:
            lookup["apply_link"] = apply_link
        elif title and company:
            lookup["title"] = title
            lookup["company"] = company

        saved_job, created = SavedJob.objects.get_or_create(defaults=defaults, **lookup)
        serializer = SavedJobSerializer(saved_job)
        return Response(
            {"message": "Job saved." if created else "Already saved.", "data": serializer.data},
            status=status.HTTP_201_CREATED if created else status.HTTP_200_OK,
        )

    def delete(self, request):
        user_id = request.query_params.get("user_id")
        apply_link = request.query_params.get("apply_link") or ""
        title = request.query_params.get("title") or ""
        company = request.query_params.get("company") or ""

        if not user_id:
            return Response({"message": "user_id is required."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response({"message": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        qs = SavedJob.objects.filter(user=user)
        if apply_link:
            qs = qs.filter(apply_link=apply_link)
        elif title and company:
            qs = qs.filter(title=title, company=company)
        else:
            return Response({"message": "apply_link or title+company required."}, status=status.HTTP_400_BAD_REQUEST)

        deleted_count, _ = qs.delete()
        if deleted_count:
            return Response({"message": "Job removed from saved."}, status=status.HTTP_200_OK)
        return Response({"message": "Saved job not found."}, status=status.HTTP_404_NOT_FOUND)


class ChatAPIView(APIView):

    def get(self, request):
        user_id = request.query_params.get("user_id")
        if not user_id:
            return Response({"message": "user_id is required."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response({"message": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        messages = ChatMessage.objects.filter(user=user).order_by("created_at")
        serializer = ChatMessageSerializer(messages, many=True)
        return Response({"results": serializer.data}, status=status.HTTP_200_OK)

    def post(self, request):
        user_id = request.data.get("user_id")
        messages = request.data.get("messages", [])

        if not user_id:
            return Response({"message": "user_id is required."}, status=status.HTTP_400_BAD_REQUEST)
        try:
            user = userModel.objects.get(id=user_id)
        except userModel.DoesNotExist:
            return Response({"message": "User not found."}, status=status.HTTP_404_NOT_FOUND)

        if not isinstance(messages, list) or not messages:
            return Response({"message": "messages must be a non-empty list."}, status=status.HTTP_400_BAD_REQUEST)

        created = []
        for msg in messages:
            if not isinstance(msg, dict):
                continue
            role = str(msg.get("role", "")).strip()
            text = str(msg.get("text", "")).strip()
            tab = str(msg.get("tab", "")).strip()
            if role not in ("user", "bot") or not text:
                continue
            obj = ChatMessage.objects.create(user=user, role=role, text=text, tab=tab)
            created.append(obj)

        serializer = ChatMessageSerializer(created, many=True)
        return Response({"results": serializer.data}, status=status.HTTP_201_CREATED)
